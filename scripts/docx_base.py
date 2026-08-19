#!/usr/bin/env python3
"""Build the Word reading edition from the Markdown knowledge base."""

from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "Cerebral_Small_Vessel_Disease_Field_Guide.docx"

DOCS = [
    ROOT / "docs" / "01_FIELD_PRIMER.md",
    ROOT / "docs" / "02_DIAGNOSTIC_CRITERIA_AND_RATING_SYSTEMS.md",
    ROOT / "docs" / "03_BIOMARKERS_AND_TOOLS.md",
    ROOT / "docs" / "04_DEBATES_HYPOTHESES_OPEN_QUESTIONS.md",
    ROOT / "docs" / "05_RESEARCH_AGENDA.md",
    ROOT / "sources" / "READING_PATH.md",
]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK_BLUE = RGBColor(32, 55, 72)
MUTED = RGBColor(90, 98, 108)
GOLD = RGBColor(143, 103, 21)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    """Keep a table row intact so labels never detach from their definitions."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_hyperlink(paragraph, text, url, color="2E74B5", underline=True):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://(?:[^()]|\([^)]*\))+)\)")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://(?:[^()]|\([^)]*\))+\)|\*[^*]+\*)")


def add_inline(paragraph, text, size=11, color=BLACK):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", size=max(8.5, size - 0.5), color=DARK_BLUE)
        elif token.startswith("["):
            label, url = LINK_RE.fullmatch(token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1
    ids = {}

    for kind, fmt, label in (("bullet", "bullet", "•"), ("decimal", "decimal", "%1.")):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), label)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        p_pr.append(ind)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(next_abs))
        num.append(abs_id)
        numbering.append(num)
        ids[kind] = next_num
        next_abs += 1
        next_num += 1
    return ids


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def clone_num_instance(doc, base_num_id):
    """Create a fresh numbering instance so a new numbered list starts at 1."""
    numbering = doc.part.numbering_part.element
    base = None
    for candidate in numbering.findall(qn("w:num")):
        if int(candidate.get(qn("w:numId"))) == int(base_num_id):
            base = candidate
            break
    if base is None:
        raise ValueError(f"Numbering instance {base_num_id} not found")
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_num = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num))
    abs_node = OxmlElement("w:abstractNumId")
    abs_node.set(qn("w:val"), abstract_id)
    num.append(abs_node)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return next_num


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (11, INK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].paragraph_format.page_break_before = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, after=2, line=1)
    run = hp.add_run("CEREBRAL SMALL VESSEL DISEASE KNOWLEDGE BASE")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B8C7D9")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Field Guide v0.1  |  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(104)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RESEARCH FIELD GUIDE")
    set_run_font(r, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Cerebral Small Vessel Disease")
    set_run_font(r, size=29, color=INK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Cerebral Amyloid Angiopathy and Brain Arteriolosclerosis")
    set_run_font(r, size=15, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("Foundations, criteria, biomarkers, mechanisms, debates, and a research agenda")
    set_run_font(r, size=11, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8, line=1.15)
    r = p.add_run("A deep-learning-oriented reading archive built from the supplied Valentina and Johanna paper collections and current authoritative literature")
    set_run_font(r, size=10.5, color=INK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(74)
    r = p.add_run("Version 0.1  |  18 August 2026")
    set_run_font(r, size=10, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Educational and research use - not patient-specific clinical guidance")
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_front_matter(doc, nums):
    p = doc.add_paragraph("How to use this guide", style="Heading 1")
    p.paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    add_inline(
        p,
        "Read the Field Primer first, then use the criteria and tool chapters as references. The debates chapter is deliberately organized around evidence for, constraints, and discriminating studies. Diagnostic rules are kept separate from claims about biological truth.",
    )

    p = doc.add_paragraph()
    add_inline(p, "Core inference chain: ", color=INK_BLUE)
    r = p.add_run("risk and susceptibility -> vessel wall/neurovascular unit -> physiology -> microscopic tissue injury -> visible imaging marker -> clinical outcome")
    set_run_font(r, size=10.5, color=INK_BLUE, bold=True)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    set_paragraph_spacing(p, before=5, after=10, line=1.2)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)

    p = doc.add_paragraph("Contents", style="Heading 2")
    chapters = [
        "1. Field primer",
        "2. Diagnostic criteria and rating systems",
        "3. Biomarkers and tools",
        "4. Debates, hypotheses, and open questions",
        "5. Research agenda",
        "6. Reading path",
        "Appendix A. Curated resource archive",
    ]
    for item in chapters:
        p = doc.add_paragraph()
        apply_num(p, nums["bullet"])
        add_inline(p, item)
        set_paragraph_spacing(p, after=4, line=1.15)

    p = doc.add_paragraph("Evidence labels", style="Heading 2")
    labels = [
        ("Definition/consensus", "standardizes language or criteria; not automatic mechanistic proof."),
        ("Neuropathology-linked", "compares an observation with tissue; sampling and interval still matter."),
        ("Longitudinal human", "supports sequence or prognosis but remains vulnerable to confounding."),
        ("Experimental", "tests mechanism in a model; human translation must be demonstrated."),
        ("Candidate biomarker", "promising but not established for all diagnostic or clinical uses."),
        ("Validated for a context of use", "tested for a stated purpose and population only."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Label"
    hdr[1].text = "Meaning"
    for label, meaning in labels:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = meaning
    format_table(table, [2500, 6860])


def format_table(table, widths):
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        prevent_row_split(row)
        for col_idx, cell in enumerate(row.cells):
            if row_idx == 0:
                shade_cell(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, after=2, line=1.1)
                for run in p.runs:
                    set_run_font(run, size=9 if len(widths) >= 4 else 9.5, color=INK_BLUE if row_idx == 0 else BLACK, bold=(row_idx == 0))
        if row_idx > 0 and row_idx % 2 == 0:
            for cell in row.cells:
                shade_cell(cell, "F8FAFC")


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def choose_widths(headers, rows):
    n = len(headers)
    lengths = []
    for i in range(n):
        vals = [headers[i]] + [r[i] if i < len(r) else "" for r in rows]
        lengths.append(max(6, min(40, max(len(v) for v in vals))))
    total_weight = sum(lengths)
    raw = [int(9360 * x / total_weight) for x in lengths]
    min_width = 1200 if n <= 4 else 900
    raw = [max(min_width, x) for x in raw]
    scale = 9360 / sum(raw)
    widths = [int(x * scale) for x in raw]
    widths[-1] += 9360 - sum(widths)
    return widths


def add_wide_table_as_records(doc, headers, rows):
    for idx, row in enumerate(rows, start=1):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=5 if idx > 1 else 1, after=2, line=1.15)
        r = p.add_run(f"Record {idx}")
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        for header, value in zip(headers, row):
            if not value:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            set_paragraph_spacing(p, after=2, line=1.12)
            r = p.add_run(f"{header}: ")
            set_run_font(r, size=9.5, color=INK_BLUE, bold=True)
            add_inline(p, value, size=9.5)


def render_markdown(doc, path, nums):
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    paragraph_buffer = []
    in_code = False
    code_lines = []
    active_decimal_num = None

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            add_inline(p, " ".join(x.strip() for x in paragraph_buffer))
            set_paragraph_spacing(p)
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            flush_paragraph()
            active_decimal_num = None
            if in_code:
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.22)
                    p.paragraph_format.right_indent = Inches(0.22)
                    set_paragraph_spacing(p, before=3, after=8, line=1.05)
                    r = p.add_run("\n".join(code_lines))
                    set_run_font(r, name="Menlo", size=8.5, color=INK_BLUE)
                    p_pr = p._p.get_or_add_pPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), LIGHT_GRAY)
                    p_pr.append(shd)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            flush_paragraph()
            active_decimal_num = None
            i += 1
            continue
        if line.startswith("#"):
            flush_paragraph()
            active_decimal_num = None
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            level = min(level, 4)
            doc.add_paragraph(title, style=f"Heading {level}")
            i += 1
            continue
        if line.startswith("| ") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1]):
            flush_paragraph()
            active_decimal_num = None
            headers = split_table_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = split_table_row(lines[i])
                if len(row) < len(headers):
                    row += [""] * (len(headers) - len(row))
                rows.append(row[: len(headers)])
                i += 1
            if len(headers) > 4:
                add_wide_table_as_records(doc, headers, rows)
            else:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for j, header in enumerate(headers):
                    table.rows[0].cells[j].text = header
                for row in rows:
                    cells = table.add_row().cells
                    for j, value in enumerate(row):
                        cells[j].text = re.sub(r"\*\*|`", "", value)
                format_table(table, choose_widths(headers, rows))
                p = doc.add_paragraph()
                set_paragraph_spacing(p, after=2, line=1)
            continue
        m_bullet = re.match(r"^\s*-\s+(.+)", line)
        m_num = re.match(r"^\s*\d+\.\s+(.+)", line)
        if m_bullet or m_num:
            flush_paragraph()
            p = doc.add_paragraph()
            if m_num:
                if active_decimal_num is None:
                    active_decimal_num = clone_num_instance(doc, nums["decimal"])
                apply_num(p, active_decimal_num)
            else:
                active_decimal_num = None
                apply_num(p, nums["bullet"])
            add_inline(p, (m_bullet or m_num).group(1))
            set_paragraph_spacing(p, after=4)
            i += 1
            continue
        if line.startswith("> "):
            flush_paragraph()
            active_decimal_num = None
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.right_indent = Inches(0.18)
            set_paragraph_spacing(p, before=4, after=8, line=1.2)
            add_inline(p, line[2:], size=11, color=INK_BLUE)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_GRAY)
            p_pr.append(shd)
            i += 1
            continue
        active_decimal_num = None
        paragraph_buffer.append(line)
        i += 1
    flush_paragraph()


def add_bibliography_appendix(doc, nums):
    doc.add_paragraph("Appendix A. Curated Resource Archive", style="Heading 1")
    p = doc.add_paragraph()
    add_inline(
        p,
        "This appendix is generated from the master bibliography. It is a curated learning set, not a completed systematic review. Essential sources are listed first; verify links and metadata before publication.",
    )
    with (ROOT / "sources" / "master_bibliography.csv").open(encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    priority_order = {"essential": 0, "high": 1, "medium": 2, "low": 3}
    rows.sort(key=lambda r: (priority_order.get(r["priority"], 9), -int(r["year"]), r["title"]))
    current = None
    bib_num = clone_num_instance(doc, nums["decimal"])
    for row in rows:
        priority = row["priority"].title()
        if priority != current:
            doc.add_paragraph(f"{priority} sources", style="Heading 2")
            current = priority
        p = doc.add_paragraph()
        apply_num(p, bib_num)
        r = p.add_run(f"{row['title']} ({row['year']}). ")
        set_run_font(r, size=10, color=BLACK, bold=True)
        r = p.add_run(f"{row['source']}; {row['evidence_type']}; {row['topic']}. ")
        set_run_font(r, size=10, color=BLACK)
        if row["doi_or_url"]:
            add_hyperlink(p, "Source", row["doi_or_url"])
            r = p.add_run(". ")
            set_run_font(r, size=10)
        r = p.add_run(row["archive_note"])
        set_run_font(r, size=9.5, color=MUTED, italic=True)
        set_paragraph_spacing(p, after=5, line=1.15)


def set_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "Cerebral Small Vessel Disease Field Guide"
    doc.core_properties.subject = "CAA, brain arteriolosclerosis, STRIVE-2, diagnostic criteria, biomarkers, and research agenda"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "cerebral small vessel disease; CAA; arteriolosclerosis; STRIVE; Boston criteria; ARTS"
    doc.settings.odd_and_even_pages_header_footer = False
    setup_page(doc)
    setup_styles(doc)
    nums = create_numbering(doc)
    add_cover(doc)
    add_front_matter(doc, nums)
    for path in DOCS:
        render_markdown(doc, path, nums)
    add_bibliography_appendix(doc, nums)
    set_update_fields(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

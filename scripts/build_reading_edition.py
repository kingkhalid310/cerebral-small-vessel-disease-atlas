#!/usr/bin/env python3
"""Build the synchronized Word reading edition from canonical repository sources."""

from __future__ import annotations

import csv
import json
import textwrap
from collections import Counter, defaultdict
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import docx_base as base


REL = Path(__file__).resolve().parents[1]
KB = REL
DATA = REL / "data"
FIG = REL / "docs" / "assets" / "figures"
OUT = REL / "downloads" / "Cerebral_Small_Vessel_Disease_Evidence_Guide_v0.5.1.docx"

PYTHON = "/Users/khalidsaifullah/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3.12"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 55, 72)
MUTED = RGBColor(90, 98, 108)
GOLD = RGBColor(143, 103, 21)
GREEN = RGBColor(35, 107, 74)
RED = RGBColor(139, 45, 45)
BLACK = RGBColor(0, 0, 0)


# Bibliographic metadata is deliberately explicit about what has and has not
# been extracted.  A blank author/identifier is preferable to an invented one.
AUTHORS_OR_GROUP = {
    "R001": "Duering et al.", "R002": "Wardlaw et al.", "R003": "Blevins et al.",
    "R004": "Greenberg", "R005": "Charidimou et al.", "R006": "Charidimou et al.",
    "R007": "Rodrigues et al.", "R008": "Auriel et al.", "R009": "Cordonnier et al.",
    "R010": "Koemans et al.", "R011": "van Veluw et al.", "R012": "Perosa et al.",
    "R013": "Kozberg et al.", "R014": "van den Brink et al.", "R015": "Perosa et al.",
    "R016": "Perosa et al.", "R017": "Makkinejad et al.", "R018": "Arfanakis et al.",
    "R019": "Arfanakis et al.", "R020": "MarkVCID Consortium", "R021": "MarkVCID Consortium",
    "R022": "MarkVCID Consortium", "R023": "Gregoire et al.", "R024": "Cordonnier et al.",
    "R025": "Skrobot et al.", "R026": "Love et al.", "R027": "Staals et al.",
    "R028": "Charidimou et al.", "R029": "Reijmer et al.", "R030": "Charidimou et al.",
    "R031": "Charidimou et al.", "R033": "Neltner et al.", "R034": "Nelson et al.",
    "R035": "Llamas-Rodriguez et al.", "R036": "Pantoni", "R037": "Wardlaw et al.",
    "R045": "Kuijf et al.", "R046": "Greenberg et al.", "R047": "Charidimou et al.",
    "R048": "Greenberg et al.", "R049": "Perosa and Viswanathan", "R050": "Perosa et al.",
    "R051": "Zanon Zotin et al.", "R053": "Del Chicca et al.",
    "R054": "Liu et al.", "R055": "van Etten et al.", "R056": "Biffi",
    "R057": "Theodorou et al.", "R058": "Wu et al.", "R059": "Theodorou et al.",
    "R060": "Sleight et al.", "R061": "Panteleienko et al.",
}

IDENTIFIERS = {
    "R006": {"pmid": "35841910", "pmcid": "PMC9389452"},
    "R017": {"pmid": "34330087", "pmcid": "PMC8329541"},
    "R018": {"pmcid": "PMC11712721"},
    "R019": {"pmcid": "PMC12208800"},
    "R051": {"pmid": "38165367", "pmcid": "PMC10834125"},
    "R052": {"pmid": "38710005", "pmcid": "PMC11177590"},
    "R053": {"pmid": "42287242"},
    "R054": {"pmid": "38951718"},
    "R055": {"pmid": "42015334", "pmcid": "PMC13099593"},
    "R056": {"pmid": "36324420", "pmcid": "PMC9616336"},
    "R057": {"pmid": "36453271"},
    "R058": {"pmid": "38417710"},
    "R059": {"pmid": "41614494"},
    "R060": {"pmid": "39499872", "pmcid": "PMC11540458"},
    "R061": {"pmid": "40970476", "pmcid": "PMC12946602"},
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def write_csv(path: Path, rows: list[dict[str, str]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def build_registries() -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    sources = read_csv(KB / "sources" / "master_bibliography.csv")
    sources.extend(read_csv(DATA / "source_additions.csv"))
    sources.sort(key=lambda r: int(r["ref_id"][1:]))

    studies = read_csv(DATA / "studies.csv")
    status_by_ref = defaultdict(list)
    for study in studies:
        status_by_ref[study["ref_id"]].append(study["extraction_status"])

    enriched_sources = []
    for row in sources:
        ref_id = row["ref_id"]
        identifiers = IDENTIFIERS.get(ref_id, {})
        link = row["doi_or_url"]
        doi = link.removeprefix("https://doi.org/") if "doi.org/" in link else ""
        enriched = dict(row)
        enriched.update({
            "authors_or_group": AUTHORS_OR_GROUP.get(ref_id, "not yet extracted"),
            "doi": doi,
            "pmid": identifiers.get("pmid", ""),
            "pmcid": identifiers.get("pmcid", ""),
            "metadata_status": "partially enriched" if ref_id in AUTHORS_OR_GROUP or identifiers else "citation metadata only",
            "reviewed_on": "2026-08-18",
        })
        enriched_sources.append(enriched)

    fields = [
        "ref_id", "title", "authors_or_group", "year", "source", "evidence_type", "topic",
        "priority", "doi", "pmid", "pmcid", "doi_or_url", "metadata_status", "archive_note", "reviewed_on",
    ]
    write_csv(DATA / "source_registry.csv", enriched_sources, fields)

    screen = []
    for row in sources:
        statuses = status_by_ref.get(row["ref_id"], [])
        if "full_text_extracted" in statuses or "full_text_web_extracted" in statuses:
            extraction = "full_text_extracted"
        elif statuses:
            extraction = statuses[0]
        else:
            extraction = "citation_only"
        screen.append({
            "screen_id": "SC" + row["ref_id"][1:],
            "ref_id": row["ref_id"],
            "decision": "include",
            "reason": "In scope for foundations, diagnosis, mechanism, measurement, or mixed pathology",
            "priority": row["priority"],
            "extraction_status": extraction,
            "duplicate_of": "",
            "reviewed_on": "2026-08-18",
        })
    write_csv(
        DATA / "screening_decisions.csv",
        screen,
        ["screen_id", "ref_id", "decision", "reason", "priority", "extraction_status", "duplicate_of", "reviewed_on"],
    )
    return enriched_sources, screen


def safe_name(text: str) -> str:
    return "".join(ch.lower() if ch.isalnum() else "_" for ch in text).strip("_")


def generate_record_pages(sources: list[dict[str, str]]) -> None:
    source_by_ref = {s["ref_id"]: s for s in sources}
    studies = read_csv(DATA / "studies.csv")
    claims = read_csv(DATA / "claims.csv")
    edges = read_csv(DATA / "claim_evidence.csv")
    tools = read_csv(DATA / "tools.csv")
    study_by_id = {s["study_id"]: s for s in studies}
    edges_by_claim = defaultdict(list)
    for edge in edges:
        edges_by_claim[edge["claim_id"]].append(edge)

    for s in studies:
        ref = source_by_ref.get(s["ref_id"], {})
        lines = [
            f"# {s['study_id']}: {s['short_name']}", "",
            f"- **Citation:** {ref.get('title', s['ref_id'])} ({s['year']})",
            f"- **Reference:** {s['ref_id']}",
            f"- **Design:** {s['design']}",
            f"- **Population:** {s['population']}",
            f"- **N:** {s['n']}",
            f"- **Setting:** {s['setting']}", "",
            "## Measurement and reference", "",
            f"- **Index test or exposure:** {s['index_test_or_exposure']}",
            f"- **Reference standard:** {s['reference_standard']}",
            f"- **Outcome:** {s['outcome']}", "",
            "## Result", "", s["key_result"], "",
            "## Limitations and provisional appraisal", "",
            f"- **Major limitations:** {s['major_limitations']}",
            f"- **Framework:** {s['appraisal_framework']}",
            f"- **Provisional risk of bias:** {s['provisional_bias']}",
            f"- **Applicability:** {s['applicability']}",
            f"- **Extraction status:** {s['extraction_status']}", "",
            "> Single-reviewer educational extraction. Verify against the full report before publication or clinical use.",
        ]
        (REL / "studies" / f"{s['study_id']}_{safe_name(s['short_name'])}.md").write_text("\n".join(lines), encoding="utf-8")

    for c in claims:
        lines = [
            f"# {c['claim_id']}: {c['claim']}", "",
            f"- **Domain:** {c['domain']}",
            f"- **Confidence:** {c['confidence']}",
            f"- **Interpretation supported:** {c['interpretation_supported']}",
            f"- **Not established:** {c['not_established']}", "",
            "## Evidence relationships", "",
        ]
        for e in edges_by_claim[c["claim_id"]]:
            eid = e["evidence_id"]
            if eid in study_by_id:
                label = f"{eid}/{study_by_id[eid]['ref_id']}"
            else:
                label = eid
            lines.extend([
                f"### {e['relationship'].title()}: {label}", "",
                f"- **Population/model:** {e['population_or_model']}",
                f"- **Result:** {e['result']}",
                f"- **Weight:** {e['weight']}",
                f"- **Limitation:** {e['limitation']}",
                f"- **Extraction:** {e['extraction_status']}", "",
            ])
        lines.extend([
            "## Synthesis judgment", "",
            f"- **Dimension scores:** directness {c['directness']}/2; bias control {c['bias_control']}/2; consistency {c['consistency']}/2; precision {c['precision']}/2; transportability {c['transportability']}/2.",
            f"- **Key limitation:** {c['key_limitation']}",
            f"- **Most decisive next test:** {c['decisive_test']}",
            f"- **Reviewed:** {c['reviewed_on']}",
        ])
        (REL / "claims" / f"{c['claim_id']}_{safe_name(c['claim'])}.md").write_text("\n".join(lines), encoding="utf-8")

    for t in tools:
        lines = [
            f"# {t['tool_id']}: {t['name']}", "",
            f"- **Category:** {t['category']}",
            f"- **Intended context:** {t['intended_context']}",
            f"- **Target construct:** {t['target_construct']}",
            f"- **Required inputs:** {t['required_inputs']}",
            f"- **Output:** {t['output']}", "",
            "## Evidence and use", "",
            f"- **Development/authority:** {t['development_or_authority']}",
            f"- **Validation status:** {t['validation_status']}",
            f"- **Reference standard:** {t['reference_standard']}",
            f"- **Key strength:** {t['key_strength']}",
            f"- **Failure mode:** {t['key_failure_mode']}",
            f"- **Clinical status:** {t['clinical_status']}",
            f"- **Version/access:** {t['version_or_access']}",
            f"- **References:** {t['key_refs']}",
            f"- **Reviewed:** {t['reviewed_on']}",
        ]
        (REL / "tools" / f"{t['tool_id']}_{safe_name(t['name'])}.md").write_text("\n".join(lines), encoding="utf-8")


FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)


def wrapped(draw, xy, text, width_px, fnt, fill="#14293D", spacing=8, align="left"):
    avg = max(7, int(fnt.size * 0.52))
    lines = textwrap.wrap(text, width=max(10, width_px // avg))
    draw.multiline_text(xy, "\n".join(lines), font=fnt, fill=fill, spacing=spacing, align=align)
    return len(lines) * (fnt.size + spacing)


def rounded_box(draw, box, title, body, fill, outline="#D2DCE6", title_color="#123653"):
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    wrapped(draw, (x1 + 24, y1 + 20), title, x2 - x1 - 48, font(28, True), title_color, 5)
    wrapped(draw, (x1 + 24, y1 + 78), body, x2 - x1 - 48, font(22), "#243746", 6)


def arrow(draw, start, end, color="#4E789E", width=8):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 12), (x - 18, y + 12)], fill=color)


def canvas(title, subtitle=""):
    im = Image.new("RGB", (1600, 900), "#FFFFFF")
    d = ImageDraw.Draw(im)
    d.rectangle((0, 0, 1600, 110), fill="#163B59")
    d.text((55, 24), title, font=font(38, True), fill="white")
    if subtitle:
        d.text((58, 73), subtitle, font=font(20), fill="#D8E8F4")
    return im, d


def make_figures() -> list[dict[str, str]]:
    FIG.mkdir(parents=True, exist_ok=True)
    specs = []

    im, d = canvas("From observation to action", "Each arrow is an empirical question, not an assumption")
    labels = [
        ("Observation", "MRI, CT, fluid, clinical, or tissue signal"),
        ("Phenotype", "Distribution, morphology, multiplicity, and time"),
        ("Etiology", "Probability of CAA, arteriolosclerosis, or another cause"),
        ("Mechanism", "Clearance, BBB, inflammation, perfusion, repair"),
        ("Action", "Decision after competing risks and context"),
    ]
    xs = [45, 355, 665, 975, 1285]
    for i, ((t, b), x) in enumerate(zip(labels, xs)):
        rounded_box(d, (x, 300, x + 265, 590), t, b, "#EFF5FA" if i % 2 == 0 else "#F7F3E8")
        if i < 4:
            arrow(d, (x + 268, 445), (xs[i + 1] - 12, 445))
    d.text((410, 690), "Reliability -> biological validity -> clinical validity -> utility", font=font(30, True), fill="#7B5A12")
    specs.append(("figure_01_inference_ladder.png", im, "Five-level inference ladder separating observation, phenotype, etiology, mechanism, and clinical action."))

    im, d = canvas("The small-vessel system", "Different compartments create different patterns of vulnerability")
    nodes = [
        ("Pial/leptomeningeal artery", "CAA-prone vessel wall", (80, 250, 390, 480), "#F5E9D6"),
        ("Penetrating arteriole", "Arteriolosclerosis, pulsatility, smooth muscle", (470, 250, 780, 480), "#E9F1F8"),
        ("Capillary / NVU", "Endothelium, pericyte, astrocyte, BBB", (860, 250, 1170, 480), "#E8F4EC"),
        ("Venule / drainage", "Outflow, inflammation, perivascular spaces", (1250, 250, 1540, 480), "#F2ECF7"),
    ]
    for i, (t, b, box, fill) in enumerate(nodes):
        rounded_box(d, box, t, b, fill)
        if i < len(nodes) - 1:
            arrow(d, (box[2] + 3, 365), (nodes[i + 1][2][0] - 12, 365))
    rounded_box(d, (275, 610, 1325, 810), "Downstream tissue consequences", "Microinfarcts, lacunes, white-matter injury, microbleeds, siderosis, atrophy, network disconnection", "#F5F7F9")
    specs.append(("figure_02_vessel_compartments.png", im, "Small-vessel compartments from pial artery to venule and their downstream tissue consequences."))

    im, d = canvas("Location changes probability - not certainty", "A differential map for common lesion distributions")
    cols = [("Lobar / cortical", "CAA rises\nArteriolosclerosis remains possible\nConsider mimics", "#F5E9D6"), ("Deep / perforator", "Arteriolosclerosis rises\nCAA may coexist\nConsider embolic or other causes", "#E9F1F8"), ("Mixed distribution", "Co-pathology is likely\nSequence and lesion age matter\nAvoid forcing one label", "#F2ECF7")]
    for i, (t, b, fill) in enumerate(cols):
        rounded_box(d, (90 + i * 500, 230, 510 + i * 500, 600), t, b, fill)
    d.text((210, 690), "Interpret with: morphology + compartment + clinical presentation + acquisition + reference evidence", font=font(28, True), fill="#163B59")
    specs.append(("figure_03_distribution_differential.png", im, "Three-column differential showing lobar, deep, and mixed lesion distributions as probabilistic evidence."))

    im, d = canvas("A staged CAA progression hypothesis", "Useful framework; not a universally proven sequence")
    stages = [("1. Susceptibility", "Age, APOE, A-beta production/clearance"), ("2. Vessel deposition", "Cortical/leptomeningeal amyloid, smooth-muscle loss"), ("3. Dysfunction", "Reactivity, permeability, drainage, inflammation"), ("4. Tissue injury", "Microinfarcts, WM injury, CMB, cSS, ICH")]
    for i, (t, b) in enumerate(stages):
        x = 70 + i * 385
        rounded_box(d, (x, 260, x + 330, 590), t, b, ["#F7F3E8", "#F5E9D6", "#E9F1F8", "#F3E8E8"][i])
        if i < 3:
            arrow(d, (x + 332, 425), (x + 370, 425))
    d.arc((300, 555, 1300, 850), 10, 170, fill="#8B5A2B", width=7)
    d.text((500, 760), "Feedback: injury and inflammation may further impair clearance and vascular function", font=font(24, True), fill="#7A4D20")
    specs.append(("figure_04_caa_progression.png", im, "Four-stage CAA progression hypothesis with a feedback loop from tissue injury to clearance and vascular function."))

    im, d = canvas("Diagnostic reasoning pipeline", "Keep description, classification, and decision-making separate")
    rows = [("1", "Describe", "Sequence quality, lesion type, exact location, count, extent"), ("2", "Build differential", "CAA, arteriolosclerosis, mixed disease, mimics, other vasculopathy"), ("3", "Apply criteria", "Only within the intended population and context of use"), ("4", "Check reference limits", "Spectrum, pathology sampling, interval, blinding, thresholds"), ("5", "Report uncertainty", "What is likely, what remains possible, what would resolve it")]
    for i, (n, t, b) in enumerate(rows):
        y = 160 + i * 135
        d.ellipse((80, y, 155, y + 75), fill="#2E74B5")
        d.text((106, y + 16), n, font=font(28, True), fill="white")
        rounded_box(d, (190, y - 10, 1510, y + 95), t, b, "#F5F8FB")
    specs.append(("figure_05_diagnostic_pipeline.png", im, "Five-step diagnostic reasoning pipeline from description through uncertainty reporting."))

    im, d = canvas("Mixed pathology is a causal system", "Coexistence can mean confounding, interaction, mediation, or shared cause")
    rounded_box(d, (80, 200, 390, 410), "Shared causes", "Age, vascular risk, APOE, systemic biology", "#F7F3E8")
    rounded_box(d, (470, 165, 760, 365), "CAA", "Vascular amyloid, cortical/leptomeningeal injury", "#F5E9D6")
    rounded_box(d, (470, 500, 760, 700), "Arteriolosclerosis", "Deep/perforator and diffuse arteriolar injury", "#E9F1F8")
    rounded_box(d, (850, 165, 1140, 365), "ADNC / LATE-NC", "Plaques, tau, TDP-43, neurodegeneration", "#F2ECF7")
    rounded_box(d, (1220, 335, 1520, 560), "Cognition and stroke", "Network injury, reserve, focal events, progression", "#E8F4EC")
    for start, end in [((390, 300),(460,260)),((390,330),(460,580)),((390,250),(840,250)),((760,260),(840,260)),((760,600),(1210,470)),((1140,270),(1210,410)),((760,330),(1210,400))]:
        arrow(d, start, end, width=6)
    specs.append(("figure_06_mixed_pathology.png", im, "Causal diagram showing shared causes, CAA, arteriolosclerosis, ADNC/LATE-NC, and cognitive or stroke outcomes."))

    im, d = canvas("The evidence graph", "The document is a view; the structured relationships are the source of truth")
    rounded_box(d, (70, 260, 340, 550), "Source", "Paper, consensus, dataset, or protocol", "#E9F1F8")
    rounded_box(d, (430, 260, 700, 550), "Study / result", "Population, measurement, estimate, uncertainty", "#E8F4EC")
    rounded_box(d, (790, 260, 1060, 550), "Claim edge", "Supports, challenges, qualifies, null, context", "#F7F3E8")
    rounded_box(d, (1150, 260, 1510, 550), "Synthesis", "Confidence, boundary, decisive next test, history", "#F2ECF7")
    arrow(d, (342,405),(418,405)); arrow(d,(702,405),(778,405)); arrow(d,(1062,405),(1138,405))
    d.text((275, 685), "One paper can contain many results; one result can affect several claims differently", font=font(30, True), fill="#163B59")
    specs.append(("figure_07_evidence_graph.png", im, "Evidence graph connecting sources to study results, claim relationships, and synthesis judgments."))

    im, d = canvas("Research roadmap for diagnostic criteria", "Discovery is only the first quarter of the journey")
    stages = [("Discovery", "Candidate features\nMechanistic plausibility", "#F7F3E8"), ("Measurement", "Reliability\nQC and harmonization", "#E9F1F8"), ("Validation", "Locked external pathology\nCalibration and subgroups", "#E8F4EC"), ("Utility", "Incremental decisions\nOutcomes and harms", "#F2ECF7")]
    for i, (t, b, fill) in enumerate(stages):
        x = 80 + i * 380
        rounded_box(d, (x, 250, x + 300, 570), t, b, fill)
        if i < 3:
            arrow(d, (x + 302, 410), (x + 368, 410))
    d.text((325, 690), "Freeze -> validate -> document failures -> update transparently", font=font(32, True), fill="#7B5A12")
    specs.append(("figure_08_research_roadmap.png", im, "Four-stage research roadmap from feature discovery through clinical utility."))

    im, d = canvas("The question is the unit of maintenance", "A topic is not complete until its answer boundary is inspectable")
    items = [("Question", "Population, exposure, comparator, reference, outcome"), ("Evidence", "Independent cohorts, estimates, nulls, limitations"), ("Judgment", "Bounded, emerging, or critical gap"), ("Next test", "Falsifier and decisive feasible design")]
    for i, (title, body) in enumerate(items):
        x = 65 + i * 385
        rounded_box(d, (x, 245, x + 315, 585), title, body, ["#E9F1F8", "#E8F4EC", "#F7F3E8", "#F2ECF7"][i])
        if i < 3:
            arrow(d, (x + 318, 415), (x + 370, 415))
    d.text((250, 690), "Every update changes a traceable object - never only a paragraph", font=font(30, True), fill="#163B59")
    specs.append(("figure_09_question_lifecycle.png", im, "Four-stage research-question lifecycle from explicit review frame to decisive next test."))

    im, d = canvas("Evidence completeness is multidimensional", "Coverage, depth, independence, reference quality, and transportability can disagree")
    labels = ["Coverage", "Full-text depth", "Cohort independence", "Reference standard", "Transportability"]
    widths = [78, 58, 42, 36, 31]
    colors = ["#2E74B5", "#4E8BB8", "#7B9F78", "#B59B56", "#A85A5A"]
    for i, (label, width, color) in enumerate(zip(labels, widths, colors)):
        y = 190 + i * 120
        d.text((95, y + 16), label, font=font(25, True), fill="#243746")
        d.rounded_rectangle((440, y, 1450, y + 70), radius=20, fill="#EEF2F5")
        d.rounded_rectangle((440, y, 440 + int(10.1 * width), y + 70), radius=20, fill=color)
        d.text((1470, y + 16), f"{width}%", font=font(24, True), fill=color)
    specs.append(("figure_10_completeness_dimensions.png", im, "Illustrative completeness bars showing why source count alone does not establish evidence maturity."))

    im, d = canvas("Publication count is not replication", "Track participant lineage before judging consistency")
    rounded_box(d, (90, 285, 390, 565), "Parent cohort", "One recruitment stream and shared infrastructure", "#E9F1F8")
    children = [("Paper A", "Derivation"), ("Paper B", "Outcome association"), ("Paper C", "Secondary biomarker")]
    for i, (title, body) in enumerate(children):
        y = 145 + i * 245
        rounded_box(d, (640, y, 960, y + 185), title, body, "#F5F8FB")
        arrow(d, (392, 425), (628, y + 92), width=6)
    rounded_box(d, (1170, 285, 1510, 565), "Independent cohort", "New recruitment, locked model, external setting", "#E8F4EC")
    arrow(d, (962, 425), (1158, 425), color="#2F7D5B", width=8)
    d.text((445, 790), "Three papers from one cohort != three replications", font=font(31, True), fill="#8B2D2D")
    specs.append(("figure_11_cohort_lineage.png", im, "Cohort lineage diagram distinguishing repeated publications from independent external replication."))

    im, d = canvas("The diagnostic validation chain", "A marker is only as mature as its weakest required link")
    chain = [("Technical", "Can it be measured?"), ("Biological", "What construct does it reflect?"), ("Accuracy", "Does it classify the target?"), ("Incremental", "Does it beat a baseline?"), ("Transport", "Does it work elsewhere?"), ("Utility", "Does using it help?")]
    for i, (title, body) in enumerate(chain):
        x = 35 + i * 260
        rounded_box(d, (x, 270, x + 225, 575), title, body, "#EFF5FA" if i < 3 else "#F7F3E8")
        if i < 5:
            arrow(d, (x + 228, 425), (x + 248, 425), width=5)
    d.text((315, 690), "Repeatability does not imply specificity; association does not imply utility", font=font(28, True), fill="#7B5A12")
    specs.append(("figure_12_validation_chain.png", im, "Six-link diagnostic validation chain from technical validity through clinical utility."))

    im, d = canvas("Anatomy of a useful hypothesis", "A claim without a challenge or falsifier is an advocacy statement")
    rounded_box(d, (90, 210, 470, 470), "Evidence for", "Direct results, replication, dose or spatial pattern", "#E8F4EC")
    rounded_box(d, (90, 535, 470, 795), "Evidence against", "Counterexamples, nulls, failed transport, bias", "#F3E8E8")
    rounded_box(d, (610, 210, 990, 470), "Alternatives", "Confounding, reverse causation, measurement error", "#F7F3E8")
    rounded_box(d, (610, 535, 990, 795), "Falsifier", "A result that would materially weaken the claim", "#F2ECF7")
    rounded_box(d, (1130, 350, 1510, 650), "Decisive experiment", "The most informative feasible next test", "#E9F1F8")
    arrow(d, (995, 500), (1118, 500), width=7)
    specs.append(("figure_13_hypothesis_anatomy.png", im, "Five-part hypothesis framework including supporting evidence, challenges, alternatives, falsifier, and decisive experiment."))

    im, d = canvas("v0.5.1 synchronized knowledge architecture", "One source of truth; document, repository, and website views")
    rounded_box(d, (75, 275, 365, 585), "Structured records", "Sources, studies, questions, claims, cohorts, tools", "#E9F1F8")
    rounded_box(d, (480, 275, 770, 585), "Evidence graph", "Relationships, confidence, provenance, update history", "#E8F4EC")
    views = [("Learner", 885, 155), ("Researcher", 1220, 155), ("Methodologist", 885, 515), ("Clinician education", 1220, 515)]
    for title, x, y in views:
        rounded_box(d, (x, y, x + 285, y + 210), title, "Generated view with explicit boundaries", "#F7F3E8" if y < 300 else "#F2ECF7")
        arrow(d, (772, 430), (x - 12, y + 105), width=5)
    arrow(d, (367, 430), (468, 430))
    specs.append(("figure_14_v03_architecture.png", im, "Synchronized knowledge architecture connecting canonical chapters and evidence records to the Word edition, repository, and website."))

    outputs = []
    for filename, image, alt in specs:
        path = FIG / filename
        image.save(path, quality=95)
        outputs.append({"path": str(path), "alt": alt})
    return outputs


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, text: str, anchor: str, size=9.5):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "2E74B5"); rpr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_before(paragraph, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    return p


def set_picture_alt(doc: Document, alt: str):
    shape = doc.inline_shapes[-1]
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", alt.split(".")[0])


def add_figure(doc, figure: dict[str, str], number: int):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(figure["path"], width=Inches(6.25))
    set_picture_alt(doc, figure["alt"])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.keep_together = True
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(f"Figure {number}. ")
    base.set_run_font(r, size=9.5, color=INK, bold=True)
    r = cap.add_run(figure["alt"])
    base.set_run_font(r, size=9.5, color=MUTED, italic=True)


def setup_page_v03(doc: Document):
    base.setup_page(doc)
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    for run in hp.runs:
        run.text = ""
    r = hp.add_run("CEREBRAL SMALL VESSEL DISEASE | LIVING EVIDENCE GUIDE")
    base.set_run_font(r, size=8.5, color=MUTED, bold=True)
    fp = section.footer.paragraphs[0]
    for run in fp.runs:
        run.text = ""
    r = fp.add_run("Evidence Guide v0.5.1  |  ")
    base.set_run_font(r, size=8.5, color=MUTED)
    base.add_page_number(fp)


def add_cover_v03(doc: Document):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(94); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LIVING EVIDENCE GUIDE"); base.set_run_font(r, size=10, color=GOLD, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Cerebral Small Vessel Disease"); base.set_run_font(r, size=29, color=INK, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Cerebral Amyloid Angiopathy and Brain Arteriolosclerosis"); base.set_run_font(r, size=15, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Foundations -> explicit questions -> evidence -> contradictions -> decisive experiments"); base.set_run_font(r, size=11, color=MUTED, italic=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.left_indent = Inches(.55); p.paragraph_format.right_indent = Inches(.55)
    r = p.add_run("A question-driven learning and research system with cohort lineage, diagnostic validation profiles, falsifiable hypotheses, worked reasoning cases, and a living-update queue.")
    base.set_run_font(r, size=11, color=INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(70)
    r = p.add_run("Version 0.5.1 | 18 August 2026"); base.set_run_font(r, size=10, color=MUTED, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Educational and research use - not patient-specific clinical guidance"); base.set_run_font(r, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()
    doc.add_paragraph("[[STATIC_TOC]]")
    doc.add_page_break()


def add_release_snapshot(doc, sources, studies, claims, edges, tools, questions, cohorts, profiles, hypotheses):
    doc.add_paragraph("Release snapshot", style="Heading 1")
    p = doc.add_paragraph()
    base.add_inline(p, "Version 0.5.1 is a synchronized question-driven learning and evidence system, not a completed systematic review. The Word edition, repository records, and website are generated from the same maintained chapters and evidence tables.")
    rows = [
        ("Scholarly sources", str(len(sources)), "Curated; citation-only and extracted records are distinguishable"),
        ("Structured pivotal studies", str(len(studies)), "Design, population, result, limitations, and provisional appraisal"),
        ("Living claims", str(len(claims)), "Normalized confidence and explicit claim boundaries"),
        ("Evidence relationships", str(len(edges)), "Support, challenge, qualification, context, null, or method"),
        ("Criteria/tools", str(len(tools)), "Context of use, inputs, outputs, validation, and failure modes"),
        ("Answerable research questions", str(len(questions)), "Bounded, emerging, or critical-gap status"),
        ("Cohort lineage records", str(len(cohorts)), "Reuse and independence are explicit"),
        ("Diagnostic evidence profiles", str(len(profiles)), "Seven-link validation chain"),
        ("Falsifiable hypotheses", str(len(hypotheses)), "For, against, alternatives, falsifier, and decisive experiment"),
        ("Original figures", "14", "Designed for causal, diagnostic, and evidence-method understanding"),
    ]
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
    for i, h in enumerate(["Object", "Count", "What changed"]): table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row): cells[i].text = value
    base.format_table(table, [2200, 1000, 6160])
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Interpretive warning. "); base.set_run_font(r, size=10.5, color=RED, bold=True)
    r = p.add_run("A polished criterion or biomarker can still be population-dependent. Always inspect the intended context, reference standard, spectrum, and external validation.")
    base.set_run_font(r, size=10.5, color=INK)


def add_claim_ledger(doc, sources, studies, claims, edges, nums):
    source_by_ref = {s["ref_id"]: s for s in sources}
    study_by_id = {s["study_id"]: s for s in studies}
    by_claim = defaultdict(list)
    for e in edges: by_claim[e["claim_id"]].append(e)
    doc.add_paragraph("Appendix H. Living Claim-Evidence Ledger", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "Each card states the narrow interpretation supported, what is not established, and the evidence that changes confidence. Ratings are provisional single-reviewer judgments.")
    for c in claims:
        doc.add_paragraph(f"{c['claim_id']} | {c['claim']}", style="Heading 3")
        table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"; table.rows[0].cells[1].text = "Current judgment"
        for label, value in [
            ("Confidence", c["confidence"].title()),
            ("Supported interpretation", c["interpretation_supported"]),
            ("Not established", c["not_established"]),
            ("Key limitation", c["key_limitation"]),
            ("Decisive next test", c["decisive_test"]),
        ]:
            cells = table.add_row().cells; cells[0].text = label; cells[1].text = value
        base.format_table(table, [2200, 7160])
        p = doc.add_paragraph()
        r = p.add_run("Evidence trace"); base.set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        for e in by_claim[c["claim_id"]]:
            p = doc.add_paragraph(); base.apply_num(p, nums["bullet"])
            eid = e["evidence_id"]
            ref_id = study_by_id[eid]["ref_id"] if eid in study_by_id else eid
            r = p.add_run(f"{e['relationship'].upper()} | "); base.set_run_font(r, size=9.5, color=GREEN if e["relationship"] == "supports" else RED if e["relationship"] == "challenges" else DARK_BLUE, bold=True)
            src = source_by_ref.get(ref_id)
            if src and src.get("doi_or_url"):
                base.add_hyperlink(p, f"{eid}/{ref_id}" if eid != ref_id else ref_id, src["doi_or_url"])
            else:
                r = p.add_run(eid); base.set_run_font(r, size=9.5, color=INK, bold=True)
            r = p.add_run(f" - {e['result']} Limitation: {e['limitation']}")
            base.set_run_font(r, size=9.5, color=BLACK)
            base.set_paragraph_spacing(p, after=3, line=1.15)


def add_study_atlas(doc, sources, studies):
    source_by_ref = {s["ref_id"]: s for s in sources}
    doc.add_paragraph("Appendix I. Pivotal Study Atlas", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "These records distinguish direct full-text extraction from abstract or citation-level coverage. Appraisals are provisional and should be duplicated before publication.")
    for s in studies:
        doc.add_paragraph(f"{s['study_id']} | {s['short_name']}", style="Heading 3")
        src = source_by_ref.get(s["ref_id"], {})
        p = doc.add_paragraph()
        r = p.add_run(f"{s['design']} | N={s['n']} | {s['setting']}"); base.set_run_font(r, size=9.5, color=MUTED, italic=True)
        for label, value in [
            ("Population", s["population"]), ("Index/exposure", s["index_test_or_exposure"]),
            ("Reference", s["reference_standard"]), ("Key result", s["key_result"]),
            ("Limitations", s["major_limitations"]), ("Appraisal", f"{s['appraisal_framework']}; {s['provisional_bias']}"),
            ("Applicability", s["applicability"]), ("Extraction", s["extraction_status"]),
        ]:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.18)
            r = p.add_run(f"{label}: "); base.set_run_font(r, size=9.5, color=INK, bold=True)
            r = p.add_run(value); base.set_run_font(r, size=9.5, color=BLACK)
            base.set_paragraph_spacing(p, after=2, line=1.12)
        if src.get("doi_or_url"):
            r = p.add_run(" | "); base.set_run_font(r, size=9.5, color=MUTED)
            base.add_hyperlink(p, f"Open source record {s['ref_id']}", src["doi_or_url"])


def add_tool_registry(doc, sources, tools):
    source_by_ref = {s["ref_id"]: s for s in sources}
    doc.add_paragraph("Appendix J. Criteria, Biomarker, and Tool Registry", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "A tool is useful only for a stated job. This registry separates description, diagnosis, burden, prognosis, technical measurement, and research enrichment.")
    for t in tools:
        doc.add_paragraph(f"{t['tool_id']} | {t['name']}", style="Heading 3")
        table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"; table.rows[0].cells[1].text = "Record"
        for label, value in [
            ("Job", t["intended_context"]), ("Target", t["target_construct"]), ("Inputs", t["required_inputs"]),
            ("Output", t["output"]), ("Validation", t["validation_status"]), ("Reference", t["reference_standard"]),
            ("Strength", t["key_strength"]), ("Failure mode", t["key_failure_mode"]), ("Status", t["clinical_status"]),
        ]:
            cells = table.add_row().cells; cells[0].text = label; cells[1].text = value
        base.format_table(table, [1900, 7460])
        p = doc.add_paragraph(); r = p.add_run("Key sources: "); base.set_run_font(r, size=9.5, color=INK, bold=True)
        for i, ref_id in enumerate(t["key_refs"].split(";")):
            ref_id = ref_id.strip(); src = source_by_ref.get(ref_id)
            if i: r = p.add_run("; "); base.set_run_font(r, size=9.5)
            if src and src.get("doi_or_url"): base.add_hyperlink(p, ref_id, src["doi_or_url"])
            else:
                r = p.add_run(ref_id); base.set_run_font(r, size=9.5)


def add_question_atlas(doc, questions):
    doc.add_paragraph("Appendix A. Research Question and Completeness Atlas", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "The question is the maintained unit. Each record states the current answer boundary, the missing piece, and the most decisive feasible design. Critical gap means important and unresolved - not unimportant.")
    current_domain = None
    for q in questions:
        if q["domain"] != current_domain:
            current_domain = q["domain"]
            doc.add_paragraph(current_domain.title(), style="Heading 2")
        doc.add_paragraph(f"{q['question_id']} | {q['question']}", style="Heading 3")
        p = doc.add_paragraph()
        r = p.add_run(f"{q['evidence_status'].replace('_', ' ').upper()} | {q['priority'].upper()}"); base.set_run_font(r, size=9.5, color=RED if q["evidence_status"] == "critical_gap" else GOLD if q["evidence_status"] == "emerging" else GREEN, bold=True)
        for label, value in [("Current answer", q["current_answer"]), ("Missing piece", q["critical_missing_piece"]), ("Decisive design", q["decisive_design"])]:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.18)
            r = p.add_run(f"{label}: "); base.set_run_font(r, size=9.5, color=INK, bold=True)
            r = p.add_run(value); base.set_run_font(r, size=9.5, color=BLACK)
            base.set_paragraph_spacing(p, after=2, line=1.12)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.18)
        r = p.add_run("Frame: "); base.set_run_font(r, size=9, color=INK, bold=True)
        r = p.add_run(f"Population={q['population']}; index/exposure={q['index_or_exposure']}; reference={q['reference_standard']}; outcome={q['outcome']}")
        base.set_run_font(r, size=9, color=MUTED, italic=True)


def add_hypothesis_atlas(doc, hypotheses):
    doc.add_paragraph("Appendix B. Contradiction and Falsification Atlas", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "A useful hypothesis survives contact with its best counterevidence. These records prevent narrative enthusiasm from replacing a testable scientific claim.")
    for h in hypotheses:
        doc.add_paragraph(f"{h['hypothesis_id']} | {h['hypothesis']}", style="Heading 3")
        table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
        table.rows[0].cells[0].text = "Lens"; table.rows[0].cells[1].text = "Current record"
        for label, value in [("Confidence", h["current_confidence"].title()), ("Strongest evidence for", h["strongest_for"]), ("Strongest evidence against", h["strongest_against"]), ("Alternative explanations", h["alternative_explanations"]), ("Falsifier", h["falsifier"]), ("Decisive experiment", h["decisive_experiment"])]:
            cells = table.add_row().cells; cells[0].text = label; cells[1].text = value
        base.format_table(table, [2200, 7160])


def add_cohort_lineage(doc, cohorts, study_cohorts):
    links = defaultdict(list)
    for row in study_cohorts: links[row["cohort_id"]].append(row)
    doc.add_paragraph("Appendix C. Cohort Lineage and Independence", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "Multiple publications from one recruitment stream do not constitute independent replication. Use this atlas before counting consistency across papers.")
    for c in cohorts:
        doc.add_paragraph(f"{c['cohort_id']} | {c['name']}", style="Heading 3")
        for label, value in [("Recruitment", c["recruitment_context"]), ("Pathology", c["pathology_availability"]), ("Known reuse", c["known_reuse"]), ("Bias relevance", c["bias_relevance"]), ("Independence", c["independence_notes"])]:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.18)
            r = p.add_run(f"{label}: "); base.set_run_font(r, size=9.5, color=INK, bold=True)
            r = p.add_run(value); base.set_run_font(r, size=9.5, color=BLACK)
            base.set_paragraph_spacing(p, after=2, line=1.12)
        if links[c["cohort_id"]]:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.18)
            r = p.add_run("Mapped studies: "); base.set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
            r = p.add_run("; ".join(f"{x['study_id']} ({x['role']})" for x in links[c["cohort_id"]])); base.set_run_font(r, size=9.5, color=MUTED)


def add_diagnostic_profiles(doc, profiles):
    doc.add_paragraph("Appendix D. Diagnostic Evidence Profiles", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "The validation links are separate. Technical repeatability cannot substitute for biological specificity; diagnostic accuracy cannot substitute for utility.")
    for d in profiles:
        doc.add_paragraph(f"{d['profile_id']} | {d['name']}", style="Heading 3")
        table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
        table.rows[0].cells[0].text = "Validation link"; table.rows[0].cells[1].text = "Judgment"
        for label, value in [("Context", d["context"]), ("Technical validity", d["technical_validity"]), ("Biological validity", d["biological_validity"]), ("Diagnostic accuracy", d["diagnostic_accuracy"]), ("Prognostic validity", d["prognostic_validity"]), ("Incremental value", d["incremental_value"]), ("Transportability", d["transportability"]), ("Clinical utility", d["clinical_utility"]), ("Failure mode", d["main_failure_mode"]), ("Next validation", d["next_validation"])]:
            cells = table.add_row().cells; cells[0].text = label; cells[1].text = value
        base.format_table(table, [2200, 7160])


def add_differential_matrix(doc, rows):
    doc.add_paragraph("Appendix E. Phenotype Differential Matrix", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "The first column is an observation, not a diagnosis. Every row lists important alternatives and the evidence most capable of resolving the etiologic question.")
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
    for i, h in enumerate(["Observation", "Raises probability", "Alternatives / discriminators", "Common error"]): table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["observation"]
        cells[1].text = row["raises_probability_of"]
        cells[2].text = f"Alternatives: {row['important_alternatives']} Discriminators: {row['discriminators']}"
        cells[3].text = row["common_error"]
    base.format_table(table, [1800, 1500, 3960, 2100])


def add_worked_cases(doc, cases):
    doc.add_paragraph("Appendix F. Worked Reasoning Cases", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "These educational cases demonstrate how to use the evidence map without converting it into patient-specific clinical decision support.")
    for c in cases:
        doc.add_paragraph(f"{c['case_id']} | {c['title']}", style="Heading 3")
        for label, value in [("Scenario", c["scenario"]), ("Reasoning", c["reasoning"]), ("Do not conclude", c["do_not_conclude"]), ("Records to open", c["records_to_open"])]:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.18)
            r = p.add_run(f"{label}: "); base.set_run_font(r, size=9.5, color=RED if label == "Do not conclude" else INK, bold=True)
            r = p.add_run(value); base.set_run_font(r, size=9.5, color=BLACK)


def add_update_queue(doc, queue):
    doc.add_paragraph("Appendix G. Living Update Queue", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "The queue converts evidence gaps into reviewable work. P0 items protect central conclusions; P1 items strengthen important domains; P2 items improve publication and website readiness.")
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
    for i, h in enumerate(["Priority", "Object", "Task", "Done when"]): table.rows[0].cells[i].text = h
    for row in queue:
        cells = table.add_row().cells
        cells[0].text = row["priority"]; cells[1].text = row["object_id"]; cells[2].text = row["task"]; cells[3].text = row["done_when"]
    base.format_table(table, [900, 1100, 4200, 3160])


def add_bibliography(doc, sources, nums):
    doc.add_paragraph("Appendix K. Curated Source Registry", style="Heading 1")
    p = doc.add_paragraph(); base.add_inline(p, "This is a curated source map, not a completed systematic review. Consult the screening and extraction-status files before using it as evidence of completeness.")
    order = {"essential": 0, "high": 1, "medium": 2, "low": 3}
    rows = sorted(sources, key=lambda r: (order.get(r["priority"], 9), -int(r["year"]), r["title"]))
    current = None; num_id = base.clone_num_instance(doc, nums["decimal"])
    for row in rows:
        if row["priority"] != current:
            current = row["priority"]
            doc.add_paragraph(f"{current.title()} sources", style="Heading 2")
        p = doc.add_paragraph(); base.apply_num(p, num_id)
        r = p.add_run(f"[{row['ref_id']}] {row['title']} ({row['year']}). "); base.set_run_font(r, size=9.5, bold=True)
        r = p.add_run(f"{row['source']}; {row['evidence_type']}. "); base.set_run_font(r, size=9.5)
        if row.get("doi_or_url"): base.add_hyperlink(p, "Source", row["doi_or_url"])
        r = p.add_run(f". {row['archive_note']}"); base.set_run_font(r, size=9, color=MUTED, italic=True)
        base.set_paragraph_spacing(p, after=4, line=1.12)


def condense_hierarchy(doc: Document):
    for p in doc.paragraphs:
        if p.style.name in {"Heading 3", "Heading 4"}:
            level = p.style.name
            p.style = doc.styles["Normal"]
            p.paragraph_format.space_before = Pt(8 if level == "Heading 3" else 6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                base.set_run_font(r, size=11 if level == "Heading 3" else 10.5, color=DARK_BLUE, bold=True)


def materialize_static_toc(doc: Document):
    placeholder = next(p for p in doc.paragraphs if p.text == "[[STATIC_TOC]]")
    headings = [p for p in doc.paragraphs if p.style.name in {"Heading 1", "Heading 2"} and p.text.strip()]
    toc = insert_before(placeholder, "Heading 1")
    toc.add_run("Contents")
    add_bookmark(toc, "Contents", 1)
    bookmark_id = 10
    for idx, h in enumerate(headings, start=1):
        if h is toc or h.text == "Contents":
            continue
        anchor = f"section_{idx:03d}"
        add_bookmark(h, anchor, bookmark_id); bookmark_id += 1
        if h.style.name == "Heading 1":
            entry = insert_before(placeholder)
            entry.paragraph_format.space_after = Pt(4)
            add_internal_link(entry, h.text, anchor, size=10)
        if h.style.name == "Heading 1":
            r = h.add_run("   "); base.set_run_font(r, size=8)
            add_internal_link(h, "Back to contents", "Contents", size=8)
    placeholder._element.getparent().remove(placeholder._element)
    if doc.paragraphs:
        add_bookmark(doc.paragraphs[0], "Top", bookmark_id)


def build_docx(sources, figures):
    studies = read_csv(DATA / "studies.csv")
    claims = read_csv(DATA / "claims.csv")
    edges = read_csv(DATA / "claim_evidence.csv")
    tools = read_csv(DATA / "tools.csv")
    questions = read_csv(DATA / "research_questions.csv")
    cohorts = read_csv(DATA / "cohorts.csv")
    study_cohorts = read_csv(DATA / "study_cohorts.csv")
    profiles = read_csv(DATA / "diagnostic_profiles.csv")
    hypotheses = read_csv(DATA / "hypotheses.csv")
    differential = read_csv(DATA / "differential_matrix.csv")
    cases = read_csv(DATA / "use_cases.csv")
    queue = read_csv(DATA / "update_queue.csv")
    doc = Document()
    doc.core_properties.title = "Cerebral Small Vessel Disease Question-Driven Evidence Guide v0.5.1"
    doc.core_properties.subject = "CAA, brain arteriolosclerosis, evidence completeness, contradictions, diagnostic validation, and research design"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "cerebral small vessel disease; CAA; arteriolosclerosis; research questions; cohort lineage; falsification; Boston criteria; ARTS"
    setup_page_v03(doc)
    base.setup_styles(doc)
    nums = base.create_numbering(doc)
    add_cover_v03(doc)
    add_release_snapshot(doc, sources, studies, claims, edges, tools, questions, cohorts, profiles, hypotheses)
    add_figure(doc, figures[13], 1)

    chapter_plan = [
        (REL / "content" / "guide" / "00_HOW_TO_THINK.md", figures[0]),
        (REL / "content" / "guide" / "01_FIELD_PRIMER_V03.md", figures[1]),
        (REL / "content" / "guide" / "02_DIAGNOSTIC_CRITERIA_AND_RATING_SYSTEMS.md", figures[4]),
        (REL / "content" / "guide" / "06_DIAGNOSTIC_TRANSPORTABILITY.md", figures[2]),
        (REL / "content" / "guide" / "03_BIOMARKERS_AND_TOOLS.md", figures[7]),
        (REL / "content" / "guide" / "04_DEBATES_HYPOTHESES_OPEN_QUESTIONS.md", figures[3]),
        (REL / "content" / "guide" / "05_RESEARCH_AGENDA_V03.md", figures[5]),
        (REL / "content" / "guide" / "07_FROM_ARCHIVE_TO_RESEARCH_PROGRAM.md", figures[6]),
        (REL / "content" / "guide" / "08_EVIDENCE_COMPLETENESS.md", figures[8]),
        (REL / "content" / "guide" / "09_CONTRADICTION_ATLAS.md", figures[12]),
        (REL / "content" / "guide" / "10_COHORT_LINEAGE.md", figures[10]),
        (REL / "content" / "guide" / "11_DIAGNOSTIC_PROFILES.md", figures[11]),
        (REL / "content" / "guide" / "12_WORKED_CASES.md", None),
        (REL / "content" / "guide" / "13_LIVING_UPDATE_WORKFLOW.md", figures[9]),
        (REL / "content" / "guide" / "14_READING_PATH.md", None),
    ]
    fig_no = 2
    for path, fig in chapter_plan:
        base.render_markdown(doc, path, nums)
        if fig:
            add_figure(doc, fig, fig_no); fig_no += 1

    add_question_atlas(doc, questions)
    add_hypothesis_atlas(doc, hypotheses)
    add_cohort_lineage(doc, cohorts, study_cohorts)
    add_diagnostic_profiles(doc, profiles)
    add_differential_matrix(doc, differential)
    add_worked_cases(doc, cases)
    add_update_queue(doc, queue)
    add_claim_ledger(doc, sources, studies, claims, edges, nums)
    add_study_atlas(doc, sources, studies)
    add_tool_registry(doc, sources, tools)
    add_bibliography(doc, sources, nums)
    condense_hierarchy(doc)
    materialize_static_toc(doc)
    base.set_update_fields(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


def write_summary(sources, screen, figures, out):
    studies = read_csv(DATA / "studies.csv")
    claims = read_csv(DATA / "claims.csv")
    edges = read_csv(DATA / "claim_evidence.csv")
    tools = read_csv(DATA / "tools.csv")
    summary = {
        "release": "v0.5.1",
        "date": "2026-08-18",
        "sources": len(sources),
        "screening_records": len(screen),
        "structured_studies": len(studies),
        "claims": len(claims),
        "claim_evidence_edges": len(edges),
        "tools": len(tools),
        "research_questions": len(read_csv(DATA / "research_questions.csv")),
        "cohorts": len(read_csv(DATA / "cohorts.csv")),
        "diagnostic_profiles": len(read_csv(DATA / "diagnostic_profiles.csv")),
        "hypotheses": len(read_csv(DATA / "hypotheses.csv")),
        "figures": len(figures),
        "confidence_counts": Counter(c["confidence"] for c in claims),
        "relationship_counts": Counter(e["relationship"] for e in edges),
        "extraction_counts": Counter(s["extraction_status"] for s in studies),
        "docx": out.relative_to(REL).as_posix(),
        "caveat": "Curated scoping evidence map; not a completed systematic review; appraisals are provisional single-reviewer judgments.",
    }
    serializable = {k: dict(v) if isinstance(v, Counter) else v for k, v in summary.items()}
    (REL / "release").mkdir(parents=True, exist_ok=True)
    (REL / "release" / "reading_edition_summary.json").write_text(json.dumps(serializable, indent=2), encoding="utf-8")
    return serializable


def main():
    sources = read_csv(DATA / "source_registry.csv")
    screen = read_csv(DATA / "screening_decisions.csv")
    figures = make_figures()
    out = build_docx(sources, figures)
    summary = write_summary(sources, screen, figures, out)
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
